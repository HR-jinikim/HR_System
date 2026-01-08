import streamlit as st
import pandas as pd
from datetime import datetime
import io
import time
import base64 
from docx import Document as DocxDocument
from docx import Document
import PyPDF2

# OpenAI 라이브러리 체크
try:
    from openai import OpenAI
    openai_installed = True
except ImportError:
    openai_installed = False

# 1. 페이지 설정
st.set_page_config(page_title="채용 서포트 시스템", layout="wide")

# --- 상태 관리 ---
if 'ai_expanded' not in st.session_state:
    st.session_state.ai_expanded = False 
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = "" 

# --- 파일 텍스트 추출 함수 ---
def extract_text_from_file(uploaded_file):
    text = ""
    try:
        uploaded_file.seek(0)
        if uploaded_file.name.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                text += page.extract_text()
        elif uploaded_file.name.endswith('.docx'):
            doc = DocxDocument(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif uploaded_file.name.lower().endswith(('.png', '.jpg', '.jpeg')):
            text = "이미지 파일입니다. (텍스트 추출 불가, 뷰어로 확인 요망)"
        else:
            text = "지원하지 않는 파일 형식입니다."
    except Exception as e:
        text = f"파일을 읽는 중 오류가 발생했습니다: {e}"
    return text

# --- 데이터 로드 ---
@st.cache_data
def load_data():
    try:
        df = pd.read_excel("data.xlsx").fillna("")
        df.columns = df.columns.str.strip() 
        data = {}
        for index, row in df.iterrows():
            def get_val(col):
                return row[col] if col in row else ""
            data[row['직무명']] = {
                "jd": row['JD'],
                "questions": {
                    "Level 1": get_val('Lv1') or get_val('질문_Lv1'),
                    "Level 2": get_val('Lv2') or get_val('질문_Lv2'),
                    "Level 3(Expert Track)": get_val('Lv3 (Expert Track)'),
                    "Level 3(Manager Track)": get_val('Lv3 (Manager Track)'),
                    "Level 4(Expert Track)": get_val('Lv4 (Expert Track)'),
                    "Level 4(Manager Track)": get_val('Lv4 (Manager Track)')
                }
            }
        return data
    except Exception:
        return {} 

jd_data = load_data()

# --- 워드 파일 생성 함수 ---
def create_word_file(position, level, comments, result, question, resume_summary):
    doc = Document()
    doc.add_heading('면접 결과 리포트', 0)
    doc.add_heading('1. 기본 정보', level=1)
    doc.add_paragraph(f'면접 일시: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph(f'지원 포지션: {position}')
    doc.add_heading('2. 역량 평가', level=1)
    doc.add_paragraph(f'평가 레벨: {level}')
    doc.add_paragraph(f'질문 내용: {question}')
    doc.add_heading('3. 면접관 코멘트', level=1)
    doc.add_paragraph(comments)
    doc.add_heading('4. 종합 결과', level=1)
    doc.add_paragraph(f'채용 추천 여부: {result}')
    
    if resume_summary:
        doc.add_heading('5. (참고) 이력서 요약', level=1)
        doc.add_paragraph(resume_summary[:500] + "...")

    bio = io.BytesIO()
    doc.save(bio)
    return bio

# --- 메인 화면 시작 ---

st.title("🤝 면접 서포트 어시스턴트")
st.markdown("---")

# 직무 선택
if jd_data:
    pos_list = list(jd_data.keys())
    selected_position = st.selectbox("**[현재 채용 중인 포지션]** 예정된 면접 포지션을 선택하세요:", pos_list)
else:
    selected_position = "데이터 없음"
    st.error("엑셀 데이터를 불러오지 못했습니다.")

st.markdown("---")

# [상단 구역] JD와 평가표
col1, col2 = st.columns([1, 1]) 

# [1구역: 왼쪽] JD
with col1:
    st.info(f"📋 {selected_position} JD")
    with st.expander("🔽 JD 상세 내용 보기 (Click)", expanded=True):
        if jd_data and selected_position in jd_data:
            st.markdown(str(jd_data[selected_position]["jd"]).replace("\n", "  \n"))
        else:
            st.write("JD 내용 없음")

# [2구역: 오른쪽] 역량 레벨 체크
with col2:
    st.success("📊 역량 레벨 체크") 
    level = st.radio(
        "지원자 레벨 선택", 
        ["Level 1", "Level 2", "Level 3(Expert Track)", "Level 3(Manager Track)", "Level 4(Expert Track)", "Level 4(Manager Track)"], 
        horizontal=False,
        key="level_select"
    )
    
    current_question = "내용을 불러올 수 없습니다."
    try:
        if jd_data and selected_position in jd_data:
            q = jd_data[selected_position]['questions'].get(level, "")
            if q: current_question = q
    except:
        pass

    st.markdown("---")
    with st.expander("💡 역량 정의 및 가이드 보기 (Click)", expanded=True):
        st.warning(f"{current_question}")


# =================================================================
# 📂 [중간] 이력서 비주얼 검토 섹션
# =================================================================
st.write("")
st.divider()
st.subheader("📂 지원자 이력서 검토")

uploaded_resume = st.file_uploader("이력서 파일(PDF, Word, 이미지)을 업로드하세요.", type=["pdf", "docx", "png", "jpg", "jpeg"])

if uploaded_resume is not None:
    with st.spinner('파일 분석 중...'):
        resume_text = extract_text_from_file(uploaded_resume)
        st.session_state.resume_text = resume_text
    
    # 비주얼 미리보기
    with st.expander("👁️ 이력서 원본 미리보기 (Click)", expanded=True):
        if uploaded_resume.name.endswith('.pdf'):
            base64_pdf = base64.b64encode(uploaded_resume.getvalue()).decode('utf-8')
            pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="1000" type="application/pdf"></iframe>'
            st.markdown(pdf_display, unsafe_allow_html=True)
        elif uploaded_resume.name.lower().endswith(('.png', '.jpg', '.jpeg')):
            st.image(uploaded_resume, caption="이력서 이미지 미리보기", use_container_width=True)
        elif uploaded_resume.name.endswith('.docx'):
            st.info("ℹ️ Word 파일은 텍스트 미리보기만 지원합니다.")
            st.text_area("추출된 텍스트", resume_text, height=300)

    st.write("") 
    
    # [버튼] AI 질문 생성
    if st.button("✨ 직무 적합도 & AI 역량 평가 질문 뽑기", type="primary"):
        st.session_state.ai_expanded = True
        
        user_display_msg = "📄 이력서/JD 기반 심층 질문과 **AI 활용 역량 평가 질문**을 생성해줘."
        st.session_state.messages.append({"role": "user", "content": user_display_msg})
        
        jd_text = str(jd_data[selected_position]["jd"]) if jd_data else ""
        
        hidden_prompt = f"""
        [역할] 너는 20년차 채용 전문가야. 
        아래 [JD]와 [이력서]를 정밀 분석하여 면접 질문을 두 가지 파트로 나누어 생성해줘.
        
        [JD 내용]
        {jd_text}
        
        [이력서 내용]
        {resume_text[:4000]}
        
        [요청사항]
        ## Part 1. 직무 및 이력서 검증 (5~7개)
        ## Part 2. AI 활용 및 미래 역량 평가 (필수 3~5개)
        
        [출력 스타일]
        인사말 생략, 질문 리스트 바로 출력.
        각 질문마다 (의도: ...) 포함.
        """
        
        st.session_state['trigger_ai_analysis'] = hidden_prompt
        st.rerun()

# =================================================================
# 🤖 [중간] AI 도우미 섹션 (드롭다운)
# =================================================================
st.write("")

with st.expander("🔽 🤖 AI 도우미 & 면접 질문 생성기 (Click)", expanded=st.session_state.ai_expanded):
    
    k_col1, k_col2 = st.columns([0.8, 0.2])
    with k_col1:
        st.info("💡 이력서 내용을 바탕으로 추가 질문을 하거나, 면접 가이드를 요청해보세요.")
    with k_col2:
        api_key = st.text_input("🔑 API Key", type="password", placeholder="Key 입력", label_visibility="collapsed")

    chat_container = st.container(height=500)
    
    with chat_container:
        if not api_key: 
            st.caption("※ API Key가 없으면 체험판 모드로 동작합니다.")
        
        if "messages" not in st.session_state:
            st.session_state["messages"] = [{"role": "assistant", "content": "안녕하세요! 이력서 분석 및 AI 역량 평가 질문을 도와드립니다."}]

        for msg in st.session_state.messages:
            st.chat_message(msg["role"]).write(msg["content"])

    # AI 응답 로직
    if 'trigger_ai_analysis' in st.session_state:
        prompt_to_send = st.session_state.pop('trigger_ai_analysis')
        
        msg = ""
        if not api_key:
            time.sleep(1.5)
            msg = """📢 [체험판 결과 예시]

## Part 1. 직무 및 이력서 검증
1. **[성과 검증]** 매출 20% 성장을 이끌었던 구체적 전략은? (의도: 성과 기여도 확인)
...

## Part 2. AI 활용 및 미래 역량 평가
1. **[AI 실무 적용]** 업무 효율을 위해 생성형 AI를 활용한 경험이 있나요? (의도: AI 활용 능력)
...

(API Key를 입력하면 실제 데이터로 생성됩니다!)"""
        else:
            if openai_installed:
                try:
                    client = OpenAI(api_key=api_key)
                    messages_for_api = [{"role": "system", "content": prompt_to_send}]
                    response = client.chat.completions.create(model="gpt-3.5-turbo", messages=messages_for_api)
                    msg = response.choices[0].message.content
                except Exception as e:
                    msg = f"❌ 오류: {e}"
            else:
                msg = "❌ OpenAI 라이브러리 설치 필요"
        
        st.session_state.messages.append({"role": "assistant", "content": msg})
        with chat_container:
            st.chat_message("assistant").write(msg)
        st.rerun()

    if prompt := st.chat_input("AI에게 추가 질문 입력..."):
        st.session_state.ai_expanded = True
        st.session_state.messages.append({"role": "user", "content": prompt})
        with chat_container:
            st.chat_message("user").write(prompt)

        msg = ""
        if not api_key:
            time.sleep(1)
            msg = "📢 [체험판] 키가 입력되지 않았습니다."
        else:
            if openai_installed:
                try:
                    client = OpenAI(api_key=api_key)
                    response = client.chat.completions.create(model="gpt-3.5-turbo", messages=st.session_state.messages)
                    msg = response.choices[0].message.content
                except Exception as e:
                    msg = f"❌ 오류: {e}"
            else:
                msg = "❌ OpenAI 설치 필요"

        st.session_state.messages.append({"role": "assistant", "content": msg})
        with chat_container:
            st.chat_message("assistant").write(msg)


# =================================================================
# 👇 [하단] 면접관 코멘트 (수정됨)
# =================================================================
st.write("") 
st.divider() 

st.header("👇 면접관 코멘트 및 최종 결과")

# [수정] 안내 문구 변경 완료
comments = st.text_area(
    '면접 진행 시 자유롭게 활용하기 위해 만들어 두었습니다. "메모장"처럼 자유롭게 활용하시기 바랍니다.', 
    height=150, 
    placeholder="여기에 면접 내용을 자유롭게 기록하세요...",
    key="comments_input"
)

st.write("") 

st.subheader("최종 결과 선택")
result = st.radio(
    "채용 여부를 선택하세요:", 
    ["채용 추천 (Pass)", "보류/불합격 (Fail)"], 
    horizontal=True,
    key="result_select"
)

st.write("") 
st.write("---")

word_file = create_word_file(selected_position, level, comments, result, current_question, st.session_state.resume_text)

st.download_button(
    label="📥 결과 리포트 다운로드 (Word)",
    data=word_file.getvalue(),
    file_name=f"면접결과_{selected_position}_{datetime.now().strftime('%Y%m%d')}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    use_container_width=True,
    type="primary"
)